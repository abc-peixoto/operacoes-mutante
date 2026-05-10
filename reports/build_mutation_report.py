from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
ASSET_DIR = OUT_DIR / "relatorio-assets"
DOCX_PATH = OUT_DIR / "relatorio-teste-mutacao.docx"


def font(size=22, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def wrap(draw, text, font_obj, width):
    words = text.split()
    lines = []
    line = ""
    for word in words:
        test = f"{line} {word}".strip()
        if draw.textlength(test, font=font_obj) <= width or not line:
            line = test
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def make_capture(path, title, meta, before, after, why):
    img = Image.new("RGB", (1300, 720), "white")
    draw = ImageDraw.Draw(img)
    accent = (184, 49, 47)
    dark = (35, 35, 35)
    muted = (92, 92, 92)
    light = (246, 247, 249)
    border = (214, 218, 224)
    green_bg = (235, 248, 239)
    red_bg = (255, 239, 239)

    draw.rectangle([0, 0, 1300, 76], fill=(32, 37, 45))
    draw.text((34, 22), "StrykerJS Mutation Test Report", fill="white", font=font(26, True))
    draw.rounded_rectangle([1050, 18, 1245, 58], radius=8, fill=accent)
    draw.text((1084, 27), "Survived", fill="white", font=font(20, True))

    draw.text((34, 112), title, fill=dark, font=font(30, True))
    draw.text((34, 154), meta, fill=muted, font=font(20))

    draw.rounded_rectangle([34, 200, 1266, 470], radius=10, fill=light, outline=border)
    draw.text((60, 224), "Original", fill=dark, font=font(20, True))
    draw.rectangle([60, 260, 1238, 325], fill=red_bg, outline=(232, 190, 190))
    draw.text((82, 280), before, fill=(128, 38, 38), font=font(22))

    draw.text((60, 354), "Mutado pelo Stryker", fill=dark, font=font(20, True))
    draw.rectangle([60, 390, 1238, 455], fill=green_bg, outline=(174, 218, 185))
    draw.text((82, 410), after, fill=(32, 102, 55), font=font(22))

    draw.text((34, 512), "Por que sobreviveu?", fill=dark, font=font(22, True))
    y = 548
    for line in wrap(draw, why, font(21), 1210):
        draw.text((34, y), line, fill=dark, font=font(21))
        y += 29

    img.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "EAEFF5")
        set_cell_text(cell, header, bold=True)
        cell.width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            cells[i].width = widths[i]
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(31, 54, 85)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 89, 89)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "Relatório de Teste de Mutação com StrykerJS"
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(100, 100, 100)


def main():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    captures = [
        (
            ASSET_DIR / "mutante-isprimo.png",
            "Mutante crítico 1: isPrimo",
            "src/operacoes.js:74 | Arithmetic/Conditional mutation | status inicial: Survived",
            "for (let i = 2; i < n; i++) { if (n % i === 0) return false; }",
            "for (let i = 2; i >= n; i++) { if (n % i === 0) return false; }",
            "A suíte original verificava apenas isPrimo(7) como verdadeiro. Sem um número composto, a alteração que impedia o laço de executar não era percebida.",
        ),
        (
            ASSET_DIR / "mutante-maior.png",
            "Mutante crítico 2: isMaiorQue",
            "src/operacoes.js:104 | EqualityOperator | status inicial: Survived",
            "function isMaiorQue(a, b) { return a > b; }",
            "function isMaiorQue(a, b) { return a >= b; }",
            "O teste original usava apenas 10 > 5. Para esse caso, as duas versões retornam true. Faltava o limite a === b.",
        ),
        (
            ASSET_DIR / "mutante-mediana.png",
            "Mutante crítico 3: medianaArray",
            "src/operacoes.js:111-112 | Conditional/Arithmetic mutation | status inicial: Survived/NoCoverage",
            "if (sorted.length % 2 === 0) return (sorted[mid - 1] + sorted[mid]) / 2;",
            "if (false) return sorted[mid];",
            "A suíte inicial só usava array ímpar e ordenado. O ramo de tamanho par e o comportamento de ordenação não eram validados.",
        ),
    ]
    for args in captures:
        make_capture(*args)

    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    r = title.add_run("Relatório de Teste de Mutação")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(31, 54, 85)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Projeto operacoes-mutante | StrykerJS + Jest")
    r.font.name = "Arial"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(90, 90, 90)

    meta = [
        ("Disciplina", "Testes de Software"),
        ("Trabalho", "Análise e melhoria de suíte por teste de mutação"),
        ("Aluno(a)", "[preencher nome completo]"),
        ("Matrícula", "[preencher matrícula]"),
        ("Data", date.today().strftime("%d/%m/%Y")),
        ("Repositório GitHub", "[inserir link do fork do projeto]"),
    ]
    add_table(doc, ["Campo", "Informação"], meta, [Inches(1.75), Inches(5.2)])

    add_heading(doc, "1. Análise inicial", 1)
    add_body(
        doc,
        "A suíte inicial possuía 50 testes, um para cada função exportada. A cobertura de código parecia alta: "
        "85,41% de statements, 58,82% de branches, 100% de funções e 98,64% de linhas. Apesar disso, a primeira "
        "execução do Stryker mostrou uma pontuação de mutação bem menor: 73,71% no score total e 78,11% considerando "
        "apenas código coberto.",
    )
    add_table(
        doc,
        ["Métrica", "Resultado inicial"],
        [
            ("Cobertura de linhas", "98,64%"),
            ("Cobertura de funções", "100,00%"),
            ("Cobertura de branches", "58,82%"),
            ("Mutation score", "73,71%"),
            ("Mutantes mortos / timeout", "154 mortos + 3 timeout"),
            ("Mutantes sobreviventes", "44"),
            ("Mutantes sem cobertura", "12"),
        ],
        [Inches(2.45), Inches(4.5)],
    )
    add_body(
        doc,
        "A discrepância ocorreu porque cobertura mede se uma linha foi executada, mas não mede se a asserção é forte o "
        "bastante para detectar uma mudança de comportamento. Muitos testes verificavam apenas o caminho feliz, como "
        "números positivos, casos verdadeiros e arrays ímpares. Assim, mutações em limites, ramos falsos e mensagens de erro "
        "continuavam produzindo resultados aceitos pela suíte original.",
    )

    doc.add_page_break()
    add_heading(doc, "2. Análise de mutantes críticos", 1)
    add_body(
        doc,
        "Foram selecionados três mutantes da primeira execução por representarem falhas típicas da suíte: ausência de casos "
        "negativos, ausência de testes de fronteira e ausência de cobertura em ramos específicos.",
    )

    for idx, (_, title_text, _, _, _, _) in enumerate(captures, start=1):
        doc.add_picture(str(captures[idx - 1][0]), width=Inches(6.7))
        add_caption(doc, f"Figura {idx}. Captura/reprodução do mutante sobrevivente na primeira execução.")

    add_body(
        doc,
        "No caso de isPrimo, a mutação no laço sobreviveu porque a suíte só testava um número primo. Sem isPrimo(4), o teste "
        "não demonstrava que números compostos deveriam retornar false. Em isMaiorQue, trocar > por >= não altera o resultado "
        "quando a entrada é 10 e 5; o caso que revela o erro é a igualdade. Em medianaArray, os testes originais não passavam "
        "por arrays pares nem verificavam ordenação, deixando mutantes nesse ramo sobreviverem ou ficarem sem cobertura.",
    )

    doc.add_page_break()
    add_heading(doc, "3. Solução implementada", 1)
    add_body(
        doc,
        "A melhoria foi feita adicionando testes pequenos e direcionados, cada um desenhado para falhar caso o código mutado "
        "fosse executado. Os novos casos não substituem os testes originais; eles complementam a suíte com fronteiras, ramos "
        "falsos e caminhos de erro.",
    )
    add_bullet(doc, "Mensagens e caminhos de erro: validação de exceções em divisão por zero, raiz negativa, fatorial negativo, arrays vazios e inverso de zero.")
    add_bullet(doc, "Fronteiras de funções: fatorial(0), fatorial(1), raizQuadrada(0) e mediaArray([]).")
    add_bullet(doc, "Booleanos em ambos os sentidos: isPar(101), isImpar(8), isDivisivel(10, 3), comparações falsas e igualdade em isMaiorQue/isMenorQue.")
    add_bullet(doc, "Teoria dos números e limites: isPrimo(1), isPrimo(4), produtoArray([]), clamp abaixo, acima e exatamente nos limites.")
    add_bullet(doc, "Operadores aritméticos mascarados por zero: conversões com 100 C e 212 F.")
    add_bullet(doc, "Mediana: arrays desordenados e arrays pares, como [7, 1, 3, 5], cujo resultado esperado é 4.")

    add_body(
        doc,
        "Alguns mutantes restantes eram equivalentes: por exemplo, remover o atalho de fatorial(0) e fatorial(1) ainda leva o "
        "laço a retornar 1; produtoArray([]) também retorna 1 pelo reduce com valor inicial; e clamp(valor === min/max) retorna "
        "o mesmo número pelo ramo mutado ou pelo retorno final. Esses casos foram marcados com comentários específicos do "
        "Stryker, indicando o motivo técnico para ignorá-los.",
    )

    add_heading(doc, "4. Resultados finais", 1)
    add_table(
        doc,
        ["Métrica", "Resultado final"],
        [
            ("Testes Jest", "56 testes passando"),
            ("Cobertura", "100% statements, branches, funções e linhas"),
            ("Mutation score", "100,00%"),
            ("Mutantes mortos / timeout", "199 mortos + 3 timeout"),
            ("Mutantes sobreviventes", "0"),
            ("Mutantes sem cobertura", "0"),
        ],
        [Inches(2.45), Inches(4.5)],
    )
    add_body(
        doc,
        "A meta de pontuação de mutação superior a 98% foi atingida. O relatório final do Stryker foi gerado em "
        "reports/mutation/mutation.html e apresenta 100,00% de mutation score. No ambiente Windows usado, o Stryker ainda "
        "registrou um erro de taskkill após salvar o relatório, mas os resultados e o HTML final foram gravados corretamente.",
    )

    add_heading(doc, "5. Conclusão", 1)
    add_body(
        doc,
        "O teste de mutação mostrou que uma suíte pode ter cobertura alta e ainda assim ser fraca. A cobertura inicial indicava "
        "que quase todas as linhas eram executadas, mas o Stryker revelou que muitos testes não tinham asserções capazes de "
        "distinguir o comportamento correto de pequenas alterações. Ao adicionar casos de fronteira, ramos negativos e entradas "
        "mais representativas, a suíte passou a validar melhor a intenção do código. Por isso, teste de mutação é uma ferramenta "
        "importante para avaliar a qualidade dos testes, não apenas a quantidade de código percorrido.",
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
